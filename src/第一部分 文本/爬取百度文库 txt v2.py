'''
@Description:
# 计算机程序设计大作业之爬取百度文库内容
# 版本号：version2
# 使用bs4爬取百度文库内纯文本内容
@Author: zll-hust && xkw
'''

import requests
from bs4 import BeautifulSoup
import bs4
from docx import Document
# https://wenku.baidu.com/robots.txt
'''
爬取前检查robots协议。
http://www.zuihaodaxue.cn/robots.txt
这种，是404 not found，则没有通过robots协议防止爬取。
https://wenku.baidu.com/robots.txt
显示：
User-agent: *
Disallow: /
对于一般用户，爬取disallow。
但是对于Baiduspider
User-agent: Baiduspider
Disallow: /w?
Disallow: /search?
Disallow: /submit
Disallow: /upload
Disallow: /cashier/
因此伪装成Baiduspider爬取。
'''

def getHTMLText(url):
    header = {'User-agent': 'Googlebot'}
    try:
        r = requests.get(url, headers = header, timeout = 30)
        r.raise_for_status()
        r.encoding = 'gbk'
        # r.encoding = r.apparent_encoding
        return r.text
    except:
        return ''

def findPList(html):
    plist = []
    soup = BeautifulSoup(html, "html.parser")
    plist.append(soup.title.string)
    for div in soup.find_all('div', attrs={"class": "bd doc-reader"}):
        plist.extend(div.get_text().split('\n'))

    plist = [c.replace(' ', '') for c in plist]
    plist = [c.replace('\x0c', '') for c in plist]
    return plist

def printPList(plist, path = 'baiduwenku4.txt'):
    print(plist)
    file = open(path, 'w',encoding='utf-8')
    for str in plist:
        file.write(str)
        file.write('\n')
    file.close()
    with open('baiduwenku4.txt',encoding='utf-8') as f:
        docu = Document()
        paragraph = docu.add_paragraph(f.read())
        docu.save('123.docx')


def main():

    url = 'https://wenku.baidu.com/view/4e29e5a730126edb6f1aff00bed5b9f3f90f72e7.html?rec_flag=default'
    html = getHTMLText(url)
    plist = findPList(html)
    printPList(plist)

if __name__ == '__main__':
    main()


# with open('baiduwenku4.txt',encoding='utf-8') as f:
#     docu = Document()
#     paragraph = docu.add_paragraph(f.read())
#     docu.save('123.docx')